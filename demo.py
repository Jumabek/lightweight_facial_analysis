#!/usr/bin/env python

from pydoc import doc
import PySimpleGUI as sg
import cv2
import numpy as np
import Emotion
import numpy as np
from hand_detector.detector import  YOLO
import mediapipe as mp
from copy import deepcopy
import time
import docx
import os
from os.path import join


# INITIALIZING OBJECTS for face landmark
mp_drawing = mp.solutions.drawing_utils
mp_drawing_styles = mp.solutions.drawing_styles
mp_face_mesh = mp.solutions.face_mesh
drawing_spec = mp_drawing.DrawingSpec(thickness=1, circle_radius=1)

# parameters for text print on image
font = cv2.FONT_HERSHEY_SIMPLEX
org = (50, 50)
org_shift_right = (200, 50)
fontScale = 0.7
color = (255, 0, 0)  
thickness = 2


# using very light CV models due to hardware requirements
emotion_model = Emotion.loadModel('weights/facial_expression_model_weights.h5') 
hand_model = YOLO(weights='weights/yolo_hand_model.h5', threshold=0.8) # using YOLO model as it is light (works w/o GPU) 
face_cascade =  cv2.CascadeClassifier('weights/haarcascade_frontalface_default.xml')


def find_emotion(img): # img is detected face crop
    emotion_labels = ['angry', 'disgust', 'fear', 'happy', 'sad', 'surprise', 'neutral']
    img = cv2.resize(img, (48,48))
    img = cv2.cvtColor(img, cv2.COLOR_BGR2GRAY)    
    img = np.expand_dims(np.expand_dims(img,axis=0),axis=3)# prep img into emotion_model format (batchsize,width,height,channels)    
    emotion_predictions = emotion_model.predict(img)[0,:]    
    return  emotion_labels[np.argmax(emotion_predictions)]


def find_face(frame):# using opencv due to resource constraints (Laptop w/o gpu)
    gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
    faces = face_cascade.detectMultiScale(gray, 1.1, 10)    
    if len(faces)>0:
        return faces[0]
    else:
        return []

def log_happy_event(img_file_path):
    mydoc = docx.Document()
    msg = "Thanks for attending Industry 4.0 Workshop!"
    mydoc.add_paragraph(msg)
    mydoc.add_picture(img_file_path)
    doc_file_path = img_file_path.replace('.png','.docx')
    mydoc.add_paragraph('Cheers!')
    mydoc.save(doc_file_path)
    #os.startfile(join(os.getcwd(),doc_file_path), "print")
    return doc_file_path


def main():

    sg.theme('Black')

    display_list_column = [[sg.Image(filename='', key='-image-', )]]
    
    # define the window layout
    option_list_column = [
            [sg.Text('Please select the options',justification='left', size=(20, 1), font='Helvetica 16')],
            [sg.HSeparator()],
              [sg.Checkbox('Emotion Recognition', change_submits = True,  enable_events=True, 
              size=(20, 1), font='Helvetica 12', key='-emotion-', default=True)],              
               [sg.Checkbox('Face Detection', change_submits = True, enable_events=True,
                size=(20, 1), font='Helvetica 12', key='-faceDetection-')],
               [sg.Checkbox('Draw Landmarks', size=(20, 1), font='Helvetica 12', key='-landmark-')],
               [sg.Checkbox('Hand Gesture', size=(20, 1), font='Helvetica 12', key='-hand-')],
               ]

    layout = [
        [
            sg.Column(display_list_column,justification='right'),
            sg.VSeparator(),
            sg.Column(option_list_column,justification='left'),
        ]
    ]

    window = sg.Window('Facial Analysis with Gesture Recognition', layout).Finalize()
    

    overlap_threshold = 0.2
    happy_timer=0 # used for keeping the text in subsequent frames even if happy is not detected in next screens
    happy_face = None

    consequtive_hand_detections=0
    CONSEQ_HANDDET_THRESHOLD = 1
    
    cap = cv2.VideoCapture(0)
    #FPS = cap.get(cv2.CAP_PROP_FPS)    
    FPS = 30
    
    with mp_face_mesh.FaceMesh(min_detection_confidence=0.5, min_tracking_confidence=0.5) as face_mesh:
        while True:# ---===--- Event LOOP Read and display frames, operate the GUI --- #
            event, values = window.read(timeout=20)
            _, frame = cap.read()
            frame = cv2.flip(frame, 1)
            original = deepcopy(frame)
            region = find_face(frame)

            if  event == sg.WIN_CLOSED:
                return

            if values['-faceDetection-']:                
                if len(region)>0:
                    x,y,w,h = region 
                    frame= cv2.rectangle(frame, (x,y),(x+w,y+h), color, thickness=2)        
            
            if values['-emotion-'] and len(region)>0:                
                x,y,w,h = region
                detected_face = frame[int(y):int(y+h), int(x):int(x+w)]            
                emo = find_emotion(detected_face)
                
                if emo == 'happy':                    
                    happy_timer=5
                    margin_h = h*20//100
                    margin_w = w*20//100
                    happy_face=deepcopy(original[max(0,(y-margin_h)):min(original.shape[0],y+h+margin_h), 
                    max(0,x-margin_w):min(original.shape[1],x+w+margin_w)])
                else:
                    happy_timer-=1

                
                if emo !=None:
                    frame = cv2.putText(frame, f'{emo}', org, font, 
                    fontScale, color, thickness, cv2.LINE_AA)
 

            if happy_timer>0 and values['-hand-']:
                                           
                tl, br = hand_model.detect(image=frame) 
                if tl and br is not None: # if there is a hand detection                                    
                    if len(region)>0 and (calc_IoU(tl,br,region)<overlap_threshold):# check overlap with face detection
                        frame = cv2.rectangle(frame, (tl[0], tl[1]), (br[0], br[1]), (235, 26, 158), 2)
                        consequtive_hand_detections+=1
                else:
                    consequtive_hand_detections=0            
                if consequtive_hand_detections>=CONSEQ_HANDDET_THRESHOLD: # makes it less sensitive to noise
                    happy_timer=0
                    consequtive_hand_detections = 0
                    
                    frame = cv2.resize(frame, (640*3//2,480*3//2), interpolation= cv2.INTER_LINEAR)
                    imgbytes = cv2.imencode('.png', frame)[1].tobytes()  # ditto
                    window['-image-'].update(data=imgbytes)

                    #loggint to file and printing to printer
                    img_file = f'screenshots/{time.time()}.png'
                    cv2.imwrite(img_file,happy_face)
                    log_happy_event(img_file)
                    

                    sg.Popup("Thanks for attending the Workshop Happy moment is captured!")                
                frame = cv2.putText(frame, f'Do you want to capture this moment?', org_shift_right, font, 
                        fontScale, color, thickness, cv2.LINE_AA)
                
            if values['-landmark-']:
                frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB) # converto to RGB color space
                results = face_mesh.process(frame)
                frame = cv2.cvtColor(frame, cv2.COLOR_RGB2BGR)
                # Draw the face mesh annotations on the image.
                if results.multi_face_landmarks:
                    for face_landmarks in results.multi_face_landmarks:
                        mp_drawing.draw_landmarks(
                            image=frame,
                            landmark_list=face_landmarks,
                            connections=mp_face_mesh.FACEMESH_TESSELATION,
                            landmark_drawing_spec=None,
                            connection_drawing_spec=mp_drawing_styles
                            .get_default_face_mesh_tesselation_style())

            frame = cv2.resize(frame, (640*3//2,480*3//2), interpolation= cv2.INTER_LINEAR)
            imgbytes = cv2.imencode('.png', frame)[1].tobytes()  # ditto
            window['-image-'].update(data=imgbytes)


def calc_area(x_left,y_top,x_right,y_bottom):
    return (x_left-x_right)*(y_bottom-y_top)


def calc_IoU(tl,br,region):
    face_tl = region[:2]
    face_br = (region[0]+region[2],region[1]+region[3])
    union = calc_area(x_left=tl[0],y_top=tl[1],x_right=br[0],y_bottom=br[1])+ calc_area(x_left=face_tl[0],y_top=face_tl[1],x_right=face_br[0],y_bottom=face_br[1])
    intersection = calc_area(x_left=max(face_tl[0],tl[0]),y_top=max(face_tl[1],tl[1]),x_right=min(face_br[0],br[0]),y_bottom=min(face_br[1],br[1]))
    return intersection/union
main()
